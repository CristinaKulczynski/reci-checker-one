import re
import nltk
from nltk.corpus import stopwords
from structure import get_sections
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def detect_language(text):
    text_lower = text.lower()
    words = re.findall(r'\b\w+\b', text_lower)

    if not words:
        return None

    try:
        pt_stopwords = set(stopwords.words('portuguese'))
        en_stopwords = set(stopwords.words('english'))
        es_stopwords = set(stopwords.words('spanish'))
    except Exception:
        return None

    pt_markers = {'com', 'uma', 'ção', 'ões', 'idade', 'mente', 'ação',
                  'são', 'não', 'também', 'muito', 'anos', 'mais', 'entre', 'idosos'}
    es_markers = {'con', 'una', 'ción', 'ciones', 'edad', 'mente', 'acción',
                  'son', 'también', 'años', 'más', 'entre', 'ancianos', 'del', 'los'}
    en_markers = {'the', 'with', 'and', 'for', 'are', 'was', 'were', 'been',
                  'have', 'has', 'had', 'their', 'which', 'this', 'that', 'from', 'will'}

    pt_count = sum(1 for word in words if word in pt_stopwords)
    en_count = sum(1 for word in words if word in en_stopwords)
    es_count = sum(1 for word in words if word in es_stopwords)

    pt_marker_count = sum(2 for word in words if word in pt_markers)
    es_marker_count = sum(2 for word in words if word in es_markers)
    en_marker_count = sum(2 for word in words if word in en_markers)

    pt_char_bonus = 5 if any(char in text_lower for char in ['ã', 'õ']) else 0
    es_char_bonus = 5 if any(char in text_lower for char in [
                             'ñ', '¿', '¡']) else 0

    pt_score = pt_count + pt_marker_count + pt_char_bonus
    en_score = en_count + en_marker_count
    es_score = es_count + es_marker_count + es_char_bonus

    if 'ç' in text_lower and re.search(r'(ção|ções|associa)', text_lower):
        pt_score += 3

    if pt_score == 0 and en_score == 0 and es_score == 0:
        return None

    max_score = max(pt_score, en_score, es_score)

    if abs(pt_score - es_score) <= 2 and max_score > en_score:
        if re.search(r'\b(com|uma|são|não|mais|idosos|associadas)\b', text_lower):
            return 'portuguese'
        elif re.search(r'\b(con|una|son|más|ancianos|asociadas)\b', text_lower):
            return 'spanish'

    if pt_score == max_score:
        return 'portuguese'
    elif en_score == max_score:
        return 'english'
    return 'spanish'


def extract_titles(title_section):
    text = re.sub(r'(ARTIGO\s*ORIGINAL|ORIGINAL\s*ARTICLE|T[ÍI]TULO)', '',
                  title_section, flags=re.IGNORECASE).strip()

    text = re.split(r'\b(ORCID|\d{4}-\d{4}-\d{4}-\d{4}|Endereço|E-mail:|Submetido|Aceite)',
                    text, maxsplit=1)[0].strip()

    author_patterns = r'\b([A-Z][a-z]+\s+[A-Z][a-z]+\s+[A-Z][a-z]+\d|[A-Z][a-z]+\s+[A-Z][a-z]+\d)'
    text = re.split(author_patterns, text, maxsplit=1)[0].strip()

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    potential_titles = [line for line in lines if len(line) >= 10]

    if len(potential_titles) >= 3:
        return potential_titles[:3]

    if len(potential_titles) == 1:
        text = potential_titles[0]

        boundary_patterns = [
            (r'(Intensiva|Brasil|Brasileir[oa]s?|Unidade|Hospital|Estudo)\s+',
             r'(Mortality|Morbidity|Prevalence|Incidence|Association|Estimating|The|A|An)\b'),
            (r'(Unit|Care|Brazil|Brazilians?|Hospital|Study)\s+',
             r'(Mortalidad|Morbilidad|Prevalencia|Incidencia|Asociación|Estimando|El|La|Los|Las)\b'),
            (r'(Intensiva|Brasil|Brasileir[oa]s?|Unidade)\s+',
             r'(Mortalidad|Morbilidad|Prevalencia|Estimando|El|La)\b'),
        ]

        boundaries = [0]

        for match in re.finditer(r'[\u200b\u200c\u200d\u2060\ufeff]', text):
            pos = match.start()
            if pos not in boundaries and pos > 10:
                boundaries.append(pos)

        for ending_pattern, starting_pattern in boundary_patterns:
            pattern = f'{ending_pattern}{starting_pattern}'
            for match in re.finditer(pattern, text, re.IGNORECASE):
                ending_match = re.match(
                    ending_pattern, text[match.start():], re.IGNORECASE)
                if ending_match:
                    boundary_pos = match.start() + ending_match.end()
                    if boundary_pos not in boundaries:
                        boundaries.append(boundary_pos)

        boundaries.append(len(text))
        boundaries.sort()

        titles = []
        for i in range(len(boundaries) - 1):
            title = text[boundaries[i]:boundaries[i+1]].strip()
            if len(title) >= 30:
                titles.append(title)

        if len(titles) >= 2:
            return titles

        return [text]

    return potential_titles


def find_first_title_paragraph(document):
    """Find the first title paragraph after the TÍTULO/ARTIGO ORIGINAL marker."""
    title_found = False

    for para in document.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        if re.match(r'^(DOI:|ORCID|E-mail:|Submetido|Aceite|\d+$)', text, re.IGNORECASE):
            continue

        if re.match(r'^(TÍTULO|ARTIGO\s*ORIGINAL|ORIGINAL\s*ARTICLE|T[ÍI]TULO)$', text, re.IGNORECASE):
            title_found = True
            continue

        if title_found:
            return para

    return None


def find_title_paragraphs(document, count=3):
    """Find multiple title paragraphs after the TÍTULO/ARTIGO ORIGINAL marker."""
    title_found = False
    title_paragraphs = []

    for para in document.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        if re.match(r'^(DOI:|ORCID|E-mail:|Submetido|Aceite|\d+$)', text, re.IGNORECASE):
            continue

        if re.match(r'^(TÍTULO|ARTIGO\s*ORIGINAL|ORIGINAL\s*ARTICLE|T[ÍI]TULO)$', text, re.IGNORECASE):
            title_found = True
            continue

        if title_found and len(text) >= 10:
            title_paragraphs.append(para)
            if len(title_paragraphs) >= count:
                break

    return title_paragraphs


def is_paragraph_centered(paragraph):
    """Check if a paragraph is centered by examining the XML element."""
    try:
        pPr = paragraph._element.pPr
        if pPr is not None:
            jc_list = pPr.findall(f'.//{WORD_NAMESPACE}jc')
            for jc in jc_list:
                if jc.get(f'{WORD_NAMESPACE}val') == 'center':
                    return True
    except:
        pass

    alignment = paragraph.alignment
    format_alignment = paragraph.paragraph_format.alignment
    return alignment == WD_ALIGN_PARAGRAPH.CENTER or format_alignment == WD_ALIGN_PARAGRAPH.CENTER


def check_paragraph_bold(paragraph):
    """Check if a paragraph has bold formatting."""
    return any(run.text.strip() and run.bold for run in paragraph.runs)


def check_paragraph_italic(paragraph):
    """Check if a paragraph has italic formatting."""
    return any(run.text.strip() and run.italic for run in paragraph.runs)


def check_paragraph_font_size(paragraph, expected_size_pt=12):
    """Check if all runs in a paragraph have the expected font size."""
    wrong_sizes = [f"{run.font.size.pt:.0f}pt" 
                   for run in paragraph.runs 
                   if run.text.strip() and run.font.size is not None and run.font.size.pt != expected_size_pt]
    return len(wrong_sizes) == 0, wrong_sizes


def is_line_spacing_simple(paragraph):
    """Check if paragraph has simple (single) line spacing."""
    line_spacing = paragraph.paragraph_format.line_spacing

    if line_spacing is None:
        return True, "default"

    if isinstance(line_spacing, float):
        is_simple = line_spacing <= 1.1
        return is_simple, f"{line_spacing:.1f}"

    is_simple = line_spacing <= 240
    return is_simple, "240 (simples)" if is_simple else f"{line_spacing}"


def item_19(manuscript_text):
    description = "Com no máximo 15 palavras."
    status = "NA"
    comments = []

    try:
        sections = get_sections(manuscript_text)

        if "TÍTULO / ORIGINAL ARTICLE" not in sections:
            return {"item": 19, "description": description, "status": status,
                    "comments": "Seção de título não encontrada"}

        titles = extract_titles(sections["TÍTULO / ORIGINAL ARTICLE"])

        if not titles:
            return {"item": 19, "description": description, "status": status,
                    "comments": "Nenhum título encontrado na seção"}

        titles_exceeding_limit = []
        all_valid = True

        for i, title in enumerate(titles):
            normalized_title = re.sub(r'(\w)-(\w)', r'\1_HYPHEN_\2', title)
            words = re.findall(r'\b\w+\b', normalized_title)
            word_count = len(words)

            if word_count > 15:
                all_valid = False
                lang = detect_language(title)
                lang_str = lang if lang else "idioma desconhecido"
                titles_exceeding_limit.append(
                    f"Título {i+1} ({lang_str}): {word_count} palavras")

        if all_valid:
            status = "A"
            comments.append(
                f"Todos os títulos têm no máximo 15 palavras ({len(titles)} títulos verificados)")
        else:
            status = "NA"
            comments.append(
                f"Alguns títulos excedem o limite de 15 palavras: {'; '.join(titles_exceeding_limit)}")

    except Exception as e:
        status = "-"
        comments.append(f"Erro ao processar títulos: {str(e)}")

    return {"item": 19, "description": description, "status": status,
            "comments": "; ".join(comments) if comments else ""}


def item_20(manuscript_path):
    description = "Principal em negrito, centralizado, fonte tamanho 12, e espaçamento simples."
    status = "NA"
    comments = []

    try:
        document = docx.Document(str(manuscript_path))

        title_paragraph = find_first_title_paragraph(document)

        if not title_paragraph:
            return {"item": 20, "description": description, "status": status,
                    "comments": "Título principal não encontrado"}

        issues = []

        if not is_paragraph_centered(title_paragraph):
            issues.append("não está centralizado")

        is_simple, spacing_desc = is_line_spacing_simple(title_paragraph)
        if not is_simple:
            issues.append(f"espaçamento não é simples ({spacing_desc})")

        if not check_paragraph_bold(title_paragraph):
            issues.append("não está em negrito")

        all_correct, wrong_sizes = check_paragraph_font_size(
            title_paragraph, expected_size_pt=12)
        if not all_correct:
            issues.append(
                f"tamanho da fonte incorreto ({', '.join(set(wrong_sizes))} ao invés de 12pt)")

        if not issues:
            status = "A"
            comments.append("Título principal está corretamente formatado")
        else:
            status = "NA"
            comments.append(f"Título principal: {'; '.join(issues)}")

    except Exception as e:
        status = "-"
        comments.append(f"Erro ao verificar formatação do título: {str(e)}")

    return {"item": 20, "description": description, "status": status,
            "comments": "; ".join(comments) if comments else ""}

def item_21(manuscript_path):
    description = "Títulos secundários em itálico, centralizado, fonte tamanho 12, e espaçamento simples."
    status = "NA"
    comments = []

    try:
        document = docx.Document(str(manuscript_path))
        title_paragraphs = find_title_paragraphs(document, count=3)

        if len(title_paragraphs) < 2:
            return {"item": 21, "description": description, "status": status,
                    "comments": "Títulos secundários não encontrados (necessário pelo menos 2 títulos)"}

        secondary_titles = title_paragraphs[1:3]
        all_issues = []

        for idx, title_paragraph in enumerate(secondary_titles, start=2):
            issues = []

            if not is_paragraph_centered(title_paragraph):
                issues.append("não está centralizado")

            is_simple, spacing_desc = is_line_spacing_simple(title_paragraph)
            if not is_simple:
                issues.append(f"espaçamento não é simples ({spacing_desc})")

            if not check_paragraph_italic(title_paragraph):
                issues.append("não está em itálico")

            all_correct, wrong_sizes = check_paragraph_font_size(title_paragraph, expected_size_pt=12)
            if not all_correct:
                issues.append(f"tamanho da fonte incorreto ({', '.join(set(wrong_sizes))} ao invés de 12pt)")

            if issues:
                all_issues.append(f"Título {idx}: {'; '.join(issues)}")

        if not all_issues:
            status = "A"
            comments.append(f"Títulos secundários estão corretamente formatados ({len(secondary_titles)} títulos verificados)")
        else:
            status = "NA"
            comments.extend(all_issues)

    except Exception as e:
        status = "-"
        comments.append(f"Erro ao verificar formatação dos títulos secundários: {str(e)}")

    return {"item": 21, "description": description, "status": status,
            "comments": "; ".join(comments) if comments else ""}


def item_23(manuscript_text):
    description = "Títulos nos três idiomas (português, inglês e espanhol)."
    status = "NA"
    comments = []

    try:
        sections = get_sections(manuscript_text)

        if "TÍTULO / ORIGINAL ARTICLE" not in sections:
            return {"item": 23, "description": description, "status": status,
                    "comments": "Seção de título não encontrada"}

        titles = extract_titles(sections["TÍTULO / ORIGINAL ARTICLE"])

        if not titles:
            return {"item": 23, "description": description, "status": status,
                    "comments": "Nenhum título encontrado na seção"}

        languages_found = set()
        for title in titles:
            lang = detect_language(title)
            if lang:
                languages_found.add(lang)

        required_languages = {'portuguese', 'english', 'spanish'}
        missing_languages = required_languages - languages_found

        if not missing_languages:
            status = "A"
            comments.append(
                f"Títulos encontrados nos três idiomas ({len(titles)} títulos detectados)")
        elif len(languages_found) >= 2:
            status = "AP"
            found = ', '.join(sorted(languages_found))
            missing = ', '.join(sorted(missing_languages))
            comments.append(
                f"Títulos parcialmente aprovados. Encontrados: {found}. Faltando: {missing}")
        else:
            found = ', '.join(sorted(languages_found)
                              ) if languages_found else 'nenhum'
            comments.append(
                f"Títulos insuficientes. Idiomas detectados: {found}")

    except Exception as e:
        status = "-"
        comments.append(f"Erro ao processar títulos: {str(e)}")

    return {"item": 23, "description": description, "status": status,
            "comments": "; ".join(comments) if comments else ""}
