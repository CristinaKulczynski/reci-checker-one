from pathlib import Path
from zipfile import ZipFile
import xml.dom.minidom
import docx

CORE = "docProps/core.xml"
APP  = "docProps/app.xml"
DOC  = "word/document.xml"

def _read_entry(docx: Path, entry: str) -> str | None:
    with ZipFile(docx) as zf:
        if entry not in zf.namelist():
            return None
        with zf.open(entry) as f:
            raw = f.read().decode("utf-8", errors="replace")
            try:
                return xml.dom.minidom.parseString(raw).toprettyxml(indent="  ")
            except Exception:
                return raw

def dump_docx_xmls(prefix: str, docx_path: Path) -> None:
    core = _read_entry(docx_path, CORE) or ""
    app  = _read_entry(docx_path, APP)  or ""
    doc  = _read_entry(docx_path, DOC)  or ""

    output_dir = Path(__file__).resolve().parent.parent / "resources" / "xml"
    output_dir.mkdir(exist_ok=True, parents=True)

    (output_dir / f"{prefix}_core.xml").write_text(core, encoding="utf-8")
    (output_dir / f"{prefix}_app.xml").write_text(app, encoding="utf-8")
    (output_dir / f"{prefix}_document.xml").write_text(doc, encoding="utf-8")

def docx_to_text(prefix: str, path: Path):
    if path.suffix.lower() != ".docx" or not path.is_file():
        raise ValueError(f"Arquivo inválido: {path}")

    dump_docx_xmls(prefix, path)
    documento = docx.Document(str(path))
    partes_texto = []

    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip():
            partes_texto.append(paragrafo.text.strip())

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if celula.text.strip():
                    partes_texto.append(celula.text.strip())

    return "\n".join(partes_texto)

def docx_pages_count(app_xml_path: Path) -> int | None:
    if not app_xml_path.exists():
        return None
    try:
        tree = ET.parse(app_xml_path)
        root = tree.getroot()
        for elem in root.iter():
            if elem.tag.endswith("Pages"):
                return int(elem.text)
    except Exception:
        return None
    return None

def docx_images_count(path: Path) -> int:
    document = docx.Document(str(path))
    count = 0
    # procura imagens em parágrafos
    for paragraph in document.paragraphs:
        count += len(paragraph._element.xpath('.//w:drawing | .//w:pict'))
    # procura imagens dentro de tabelas
    for tabela in document.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                count += len(celula._element.xpath('.//w:drawing | .//w:pict'))
    return count
